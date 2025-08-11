from docx import Document
from docx.shared import RGBColor
from pathlib import Path
import uuid, json

# We will NOT use MS Word's XML comment model (complex). Instead:
# - highlight the run (by setting font color or wrap with brackets)
# - append a comment table at end of doc listing: comment id, location, suggestion, evidence
# - return JSON mapping of comments for the system to use.

def annotate_docx(input_path: str, annotations: dict, out_path: str = None):
    """
    annotations: dict mapping of { 'search_text_or_heading' : { 'comment': str, 'severity': 'High' } }
    We'll find occurrences of the search_text and wrap in [!! ... !!] and add a bold note.
    """
    doc = Document(input_path)
    out_path = out_path or (str(Path(input_path).with_name(Path(input_path).stem + "_reviewed.docx")))
    comment_list = []
    for p in doc.paragraphs:
        for key, note in annotations.items():
            if key.strip() == "":
                continue
            if key in p.text:
                # naive replace, preserve run structure if possible
                p.text = p.text.replace(key, f"[!!{key}!!]")
                comment_list.append({"location": p.text[:120], "original": key, "comment": note["comment"], "severity": note.get("severity","Medium")})
    # append comment summary at end
    doc.add_page_break()
    doc.add_paragraph("=== REVIEW COMMENTS ===")
    for c in comment_list:
        p = doc.add_paragraph()
        p.add_run(f"ID: {uuid.uuid4().hex[:8]} | Severity: {c['severity']}\n").bold = True
        p.add_run(f"Location snippet: {c['location']}\n")
        p.add_run(f"Issue: {c['original']}\n")
        p.add_run(f"Recommendation: {c['comment']}\n")
        p.add_run("-"*40 + "\n")
    doc.save(out_path)
    meta_path = out_path + ".comments.json"
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(comment_list, f, ensure_ascii=False, indent=2)
    return out_path, meta_path
